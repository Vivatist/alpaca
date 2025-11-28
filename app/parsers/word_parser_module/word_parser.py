#!/usr/bin/env python3
"""
Word Document Parser для RAG системы ALPACA

Парсер Word документов (.doc, .docx) с OCR поддержкой для отсканированных страниц.

Pipeline:
    .docx/.doc → Markitdown + Unstructured OCR → Markdown + YAML metadata

Возможности:
- Извлечение текста с сохранением структуры (заголовки, списки, таблицы)
- OCR для встроенных изображений (отсканированные документы)
- Извлечение метаданных (автор, дата создания, количество страниц)
- Генерация YAML header с метаданными
- Экспорт в Markdown формат для RAG индексации
"""

import os
import sys
import tempfile
import shutil
import subprocess
from pathlib import Path
from typing import Dict, Optional, List, TYPE_CHECKING

# Добавляем путь к базовому парсеру
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
from base_parser import BaseParser

if TYPE_CHECKING:
    from utils.file_manager import File

from markitdown import MarkItDown  # type: ignore
from unstructured.partition.auto import partition  # type: ignore
from unstructured.partition.image import partition_image  # type: ignore
from docx import Document  # type: ignore
from docx.oxml.table import CT_Tbl  # type: ignore
from docx.oxml.text.paragraph import CT_P  # type: ignore
from docx.table import _Cell, Table  # type: ignore
from docx.text.paragraph import Paragraph  # type: ignore
from PIL import Image  # type: ignore
import pytesseract  # type: ignore
import pdf2image  # type: ignore


class WordParser(BaseParser):
    """
    Парсер Word документов с поддержкой OCR
    
    Использует pipeline:
    1. Markitdown - основной парсер структуры документа
    2. python-docx - извлечение метаданных и изображений
    3. Unstructured - OCR для изображений (отсканированные страницы)
    4. YAML header - метаданные для RAG
    """
    
    def __init__(self, enable_ocr: bool = True, ocr_strategy: str = "auto"):
        """
        Инициализация парсера
        
        Args:
            enable_ocr: Включить OCR для изображений
            ocr_strategy: Стратегия OCR ('auto', 'hi_res', 'fast')
        """
        super().__init__("word-parser")
        self.enable_ocr = enable_ocr
        self.ocr_strategy = ocr_strategy
        
        # Инициализация Markitdown
        self.markitdown = MarkItDown()
    
    def parse(self, file: 'File') -> str:
        """
        Парсинг Word документа в текст
        
        Args:
            file: Объект File с информацией о файле
            
        Returns:
            str: Распарсенный текст документа (пустая строка при ошибке)
        """
        converted = False  # Инициализируем перед try для использования в finally
        file_path = file.path  # Извлекаем путь из объекта File
        file_hash = file.hash  # Извлекаем хэш из объекта File
        
        try:
            if not os.path.exists(file_path):
                self.logger.error(f"File not found | file={file_path}")
                return ""
            
            self.logger.info(f"Parsing Word document | file={file_path}")
            
            # 1. Конвертация .doc → .docx через LibreOffice если нужно
            file_ext = Path(file_path).suffix.lower()
            original_file_path = file_path
            
            if file_ext == '.doc':
                self.logger.info(f"Old .doc format detected, converting to .docx via LibreOffice")
                converted_path = self._convert_doc_to_docx(file_path)
                if converted_path:
                    file_path = converted_path
                    file_ext = '.docx'
                    converted = True
                    self.logger.info(f"Converted .doc to .docx | path={converted_path}")
                else:
                    self.logger.warning(f"Failed to convert .doc to .docx, will try direct parsing")
            
            # 2. Добавляем ОБЩИЕ метаданные (в базовом классе)
            common_metadata = self._add_common_metadata(original_file_path, file_hash)
            
            # 3. Извлечение СПЕЦИФИЧНЫХ метаданных Word через python-docx (только для .docx)
            specific_metadata = {}
            
            if file_ext == '.docx':
                # Полные метаданные только для .docx
                try:
                    specific_metadata = self._extract_word_specific_metadata(file_path)
                except Exception as e:
                    self.logger.warning(f"Failed to extract Word-specific metadata | error={type(e).__name__}: {e}")
                    # Если метаданные не извлекаются (битый .docx после конвертации),
                    # пробуем парсить напрямую через fallback
                    if converted:
                        self.logger.info(f"Converted .docx appears corrupted, using fallback parser on original .doc")
                        file_path = original_file_path
                        file_ext = '.doc'
            
            # 4. Основной парсинг через Markitdown (сначала получаем весь документ)
            markdown_content = self._parse_with_markitdown(file_path)
            
            # 5. OCR для изображений (если включено и это .docx)
            if self.enable_ocr and file_ext == '.docx':
                self.logger.info(f"Processing document with OCR | enable_ocr={self.enable_ocr}")
                
                # Извлекаем изображения
                images_info = self._extract_images(file_path)
                self.logger.info(f"Images extracted | count={len(images_info)}")
                
                if images_info:
                    self.logger.info(f"Starting OCR processing | images={len(images_info)}")
                    
                    # Получаем OCR текст для каждого изображения
                    ocr_texts = self._process_images_with_ocr_individually(images_info)
                    
                    if ocr_texts:
                        # Заменяем base64 изображения на OCR текст в правильном порядке
                        import re
                        image_pattern = r'!\[\]\(data:image/[^)]+\)'
                        
                        def replace_image(match):
                            nonlocal ocr_texts
                            if ocr_texts:
                                return ocr_texts.pop(0)  # Берем следующий OCR текст
                            return match.group(0)  # Если OCR текстов не хватает, оставляем как есть
                        
                        markdown_content = re.sub(image_pattern, replace_image, markdown_content)
                        self.logger.info(f"OCR content inserted | replaced_images={len(images_info) - len(ocr_texts)}")
                    else:
                        self.logger.warning("OCR processing returned no content")
                else:
                    self.logger.info("No images found in document for OCR")
            elif not self.enable_ocr:
                self.logger.debug("OCR disabled, skipping image processing")
            
            self.logger.info(f"Word document parsed successfully | file={original_file_path} length={len(markdown_content)}")
            
            return markdown_content
            
        except Exception as e:
            self.logger.error(f"Error parsing Word document | file={file_path} error={type(e).__name__}: {e}")
            return ""
        
        finally:
            # 9. Очистка временных файлов после конвертации
            if converted:
                try:
                    converted_dir = Path(file_path).parent
                    if converted_dir.name.startswith("alpaca_doc_convert_"):
                        shutil.rmtree(converted_dir)
                        self.logger.info(f"Cleaned up temp conversion directory | dir={converted_dir}")
                except Exception as e:
                    self.logger.warning(f"Failed to clean up temp directory | error={type(e).__name__}: {e}")
    
    def _convert_doc_to_docx(self, doc_path: str) -> Optional[str]:
        """
        Конвертация .doc в .docx через LibreOffice headless
        
        Args:
            doc_path: Путь к исходному .doc файлу
            
        Returns:
            Путь к сконвертированному .docx или None при ошибке
        """
        try:
            # Создаем временную директорию для конвертированного файла
            temp_dir = tempfile.mkdtemp(prefix="alpaca_doc_convert_")
            self.logger.info(f"Converting .doc to .docx (source={doc_path} temp_dir={temp_dir})")
            
            # Запускаем LibreOffice в headless режиме
            result = subprocess.run(
                [
                    'libreoffice',
                    '--headless',
                    '--convert-to', 'docx',
                    '--outdir', temp_dir,
                    doc_path
                ],
                capture_output=True,
                text=True,
                timeout=60
            )
            
            if result.returncode != 0:
                self.logger.error(f"LibreOffice conversion failed (returncode={result.returncode} stderr={result.stderr})")
                return None
            
            # Проверяем что файл создан
            doc_filename = Path(doc_path).stem
            converted_file = Path(temp_dir) / f"{doc_filename}.docx"
            
            if converted_file.exists():
                self.logger.info(f"Conversion successful (output={converted_file} size={converted_file.stat().st_size})")
                return str(converted_file)
            else:
                self.logger.error(f"Converted file not found (expected={converted_file})")
                return None
                
        except subprocess.TimeoutExpired:
            self.logger.error(f"LibreOffice conversion timeout (file={doc_path} limit=60s)")
            return None
        except Exception as e:
            self.logger.error(f"Conversion error (file={doc_path} error={type(e).__name__}: {e})")
            return None
    
    def _extract_word_specific_metadata(self, file_path: str) -> Dict:
        """
        Извлечение СПЕЦИФИЧНЫХ для Word метаданных
        
        Общие метаданные (file_name, file_path, file_size, etc.) добавляются в базовом классе.
        Здесь добавляем только специфичные для Word данные.
        
        Args:
            file_path: Путь к файлу
            
        Returns:
            Dict со специфичными метаданными
        """
        specific_metadata = {
            'author': '',
            'subject': '',
            'pages': 0,
            'paragraphs': 0,
            'tables': 0,
            'images': 0
        }
        
        try:
            doc = Document(file_path)
            core_props = doc.core_properties
            
            specific_metadata['author'] = core_props.author or ''
            specific_metadata['subject'] = core_props.subject or ''
            
            # Подсчет элементов
            specific_metadata['paragraphs'] = len(doc.paragraphs)
            specific_metadata['tables'] = len(doc.tables)
            
            # Приблизительное количество страниц (250 слов на страницу)
            total_words = sum(len(p.text.split()) for p in doc.paragraphs)
            specific_metadata['pages'] = max(1, total_words // 250)
            
            # Подсчет изображений
            image_count = 0
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_count += 1
            specific_metadata['images'] = image_count
            
            self.logger.debug(f"Word-specific metadata | author={specific_metadata['author']} pages={specific_metadata['pages']} paragraphs={specific_metadata['paragraphs']}")
            
        except Exception as e:
            self.logger.warning(f"Word-specific metadata extraction failed | file={file_path} error={type(e).__name__}: {e}")

        return specific_metadata
    
    def _extract_images(self, file_path: str) -> List[Dict]:
        """
        Извлечение изображений из Word документа для OCR
        
        Args:
            file_path: Путь к файлу
            
        Returns:
            List[Dict] с информацией об изображениях
        """
        images = []
        
        try:
            doc = Document(file_path)
            temp_dir = tempfile.mkdtemp(prefix="alpaca_word_images_")
            
            self.logger.info(f"Checking for images in document | file={file_path} temp_dir={temp_dir}")
            
            image_idx = 0
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_idx += 1
                    
                    try:
                        image_data = rel.target_part.blob
                        
                        if not image_data:
                            self.logger.warning(f"Empty image data | index={image_idx}")
                            continue
                        
                        # Определяем расширение по content_type
                        content_type = rel.target_part.content_type
                        ext = self._get_image_extension(content_type)
                        
                        # Сохраняем во временную директорию
                        image_path = os.path.join(temp_dir, f"image_{image_idx}{ext}")
                        with open(image_path, 'wb') as f:
                            f.write(image_data)
                        
                        # Проверяем что файл создан
                        if not os.path.exists(image_path):
                            self.logger.error(f"Failed to save image | index={image_idx} path={image_path}")
                            continue
                        
                        self.logger.debug(f"Image saved | index={image_idx} size={len(image_data)} type={content_type}")
                        
                        # КРИТИЧНО: Конвертируем WMF/EMF в PNG для OCR
                        if content_type in ('image/x-wmf', 'image/x-emf') or ext in ('.wmf', '.emf'):
                            converted_path = self._convert_wmf_to_png(image_path, image_idx, temp_dir)
                            if converted_path:
                                image_path = converted_path
                                ext = '.png'
                                self.logger.info(f"Converted WMF/EMF to PNG | index={image_idx} path={converted_path}")
                            else:
                                self.logger.warning(f"WMF/EMF conversion failed, trying PDF method | index={image_idx}")
                                # Альтернатива: конвертируем весь DOCX в PDF, затем в изображения
                                pdf_converted = self._extract_images_via_pdf(file_path, image_idx, temp_dir)
                                if pdf_converted:
                                    image_path = pdf_converted
                                    ext = '.png'
                                    self.logger.info(f"Converted via PDF method | index={image_idx} path={pdf_converted}")
                                else:
                                    self.logger.error(f"All conversion methods failed | index={image_idx}")
                                    continue
                        
                        images.append({
                            'index': image_idx,
                            'path': image_path,
                            'size': len(image_data),
                            'type': content_type
                        })
                        
                    except Exception as e:
                        self.logger.error(f"Failed to extract image | index={image_idx} error={type(e).__name__}: {e}")
            
            self.logger.info(f"Extracted {len(images)} images for OCR (file={file_path})")
            
        except Exception as e:
            self.logger.warning(f"Error extracting images (file={file_path}): {e}")
        
        return images
    
    def _get_image_extension(self, content_type: str) -> str:
        """Определение расширения файла по MIME типу"""
        extensions = {
            'image/jpeg': '.jpg',
            'image/png': '.png',
            'image/gif': '.gif',
            'image/bmp': '.bmp',
            'image/tiff': '.tiff',
            'image/webp': '.webp',
            'image/x-wmf': '.wmf',
            'image/x-emf': '.emf'
        }
        return extensions.get(content_type, '.jpg')
    
    def _extract_images_via_pdf(self, docx_path: str, image_idx: int, temp_dir: str) -> Optional[str]:
        """
        Извлечение изображений через конвертацию DOCX→PDF→PNG
        
        Используется когда WMF/EMF не могут быть конвертированы напрямую.
        
        Args:
            docx_path: Путь к DOCX файлу
            image_idx: Индекс изображения
            temp_dir: Временная директория
            
        Returns:
            Путь к PNG файлу или None если конвертация не удалась
        """
        try:
            # Конвертируем DOCX в PDF через LibreOffice
            pdf_temp_dir = tempfile.mkdtemp(prefix="alpaca_pdf_convert_")
            
            result = subprocess.run(
                ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', pdf_temp_dir, docx_path],
                capture_output=True,
                text=True,
                timeout=60
            )
            
            if result.returncode != 0:
                self.logger.error(f"LibreOffice PDF conversion failed | returncode={result.returncode}")
                return None
            
            # Находим созданный PDF
            pdf_files = list(Path(pdf_temp_dir).glob('*.pdf'))
            if not pdf_files:
                self.logger.error(f"No PDF file created | dir={pdf_temp_dir}")
                return None
            
            pdf_path = str(pdf_files[0])
            self.logger.info(f"PDF created | path={pdf_path}")
            
            # Конвертируем PDF в изображения
            from pdf2image import convert_from_path
            images = convert_from_path(pdf_path, dpi=200)
            
            if not images:
                self.logger.error(f"No images extracted from PDF")
                return None
            
            # Сохраняем первую страницу (где обычно находится изображение)
            png_path = os.path.join(temp_dir, f"image_{image_idx}_from_pdf.png")
            images[0].save(png_path, 'PNG')
            
            # Очистка временных файлов
            shutil.rmtree(pdf_temp_dir, ignore_errors=True)
            
            self.logger.info(f"Image extracted via PDF | path={png_path}")
            return png_path
            
        except Exception as e:
            self.logger.error(f"PDF extraction failed | error={type(e).__name__}: {e}")
            return None
    
    def _convert_wmf_to_png(self, wmf_path: str, image_idx: int, temp_dir: str) -> Optional[str]:
        """
        Конвертация WMF/EMF изображения в PNG для OCR
        
        WMF (Windows Metafile) не поддерживается PIL напрямую,
        но мы можем использовать ImageMagick через subprocess или
        попробовать открыть как BMP (некоторые WMF содержат BMP data)
        
        Args:
            wmf_path: Путь к WMF файлу
            image_idx: Индекс изображения
            temp_dir: Временная директория
            
        Returns:
            Путь к PNG файлу или None если конвертация не удалась
        """
        png_path = os.path.join(temp_dir, f"image_{image_idx}_converted.png")
        
        try:
            # Попытка 1: Использовать ImageMagick через subprocess
            # ImageMagick 7 использует команду 'magick' вместо 'convert'
            import subprocess
            
            self.logger.info(f"Attempting ImageMagick conversion for image {image_idx}")
            
            # Пробуем сначала ImageMagick 7 (команда 'magick')
            try:
                result = subprocess.run(
                    ['magick', wmf_path, png_path],
                    capture_output=True,
                    timeout=30,
                    text=True
                )
                self.logger.debug(f"ImageMagick result: returncode={result.returncode}")
            except FileNotFoundError as e:
                self.logger.debug(f"'magick' command not found, trying 'convert'")
                # Fallback для ImageMagick 6 (команда 'convert')
                result = subprocess.run(
                    ['convert', wmf_path, png_path],
                    capture_output=True,
                    timeout=30,
                    text=True
                )
                self.logger.debug(f"Convert result: returncode={result.returncode}")
            
            if result.returncode == 0 and os.path.exists(png_path):
                self.logger.info(f"WMF converted with ImageMagick: image {image_idx}")
                return png_path
            else:
                self.logger.warning(f"ImageMagick conversion failed with return code {result.returncode}")
        except (FileNotFoundError, subprocess.TimeoutExpired, Exception) as e:
            self.logger.warning(f"ImageMagick conversion exception: {e}")
        
        try:
            # Попытка 2: Попробовать открыть напрямую через PIL
            # Некоторые "WMF" файлы на самом деле обычные растровые изображения
            img = Image.open(wmf_path)
            img.save(png_path, 'PNG')
            
            self.logger.info(f"WMF converted with PIL: image {image_idx}")
            return png_path
        except Exception as e:
            self.logger.warning(f"Failed to convert WMF image {image_idx}: {e}")
            return None
    
    def _parse_with_markitdown(self, file_path: str) -> str:
        """
        Парсинг документа через Markitdown
        
        Args:
            file_path: Путь к файлу
            
        Returns:
            Markdown текст
        """
        try:
            result = self.markitdown.convert(file_path)
            markdown = result.text_content if hasattr(result, 'text_content') else str(result)
            
            self.logger.debug(f"Markitdown parsing complete | length={len(markdown)}")
            
            return markdown
            
        except BaseException as e:
            # Используем BaseException чтобы поймать FileConversionException
            self.logger.warning(f"Markitdown parsing failed | error={type(e).__name__}: {str(e)[:200]}")
            # Для .doc файлов пробуем альтернативные методы
            file_ext = Path(file_path).suffix.lower()
            if file_ext == '.doc':
                self.logger.info("Trying alternative .doc parsing with olefile")
                try:
                    import olefile # type: ignore
                    if olefile.isOleFile(file_path):
                        ole = olefile.OleFileIO(file_path)
                        # Извлекаем WordDocument stream
                        if ole.exists('WordDocument'):
                            word_stream = ole.openstream('WordDocument').read()
                            # Простое извлечение текста (наивный подход)
                            text = word_stream.decode('latin-1', errors='ignore')
                            # Фильтруем печатные символы
                            text = ''.join(c for c in text if c.isprintable() or c in '\n\r\t')
                            ole.close()
                            if len(text) > 100:  # Разумное количество текста
                                self.logger.info(f"Olefile parsing successful | length={len(text)}")
                                return text
                        else:
                            self.logger.warning("WordDocument stream not found in .doc file")
                except Exception as ole_error:
                    self.logger.warning(f"Olefile parsing failed | error={type(ole_error).__name__}: {ole_error}")
            
            self.logger.info(f"Using fallback parser for {file_ext} file")
            return self._fallback_parse(file_path)
    
    def _fallback_parse(self, file_path: str) -> str:
        """
        Резервный парсер через python-docx или antiword для .doc
        
        Args:
            file_path: Путь к файлу
            
        Returns:
            Простой текст документа
        """
        file_ext = Path(file_path).suffix.lower()
        
        # Для .doc файлов сначала пробуем olefile
        if file_ext == '.doc':
            self.logger.info("Fallback parser: trying olefile for .doc file")
            try:
                import olefile # type: ignore
                if olefile.isOleFile(file_path):
                    ole = olefile.OleFileIO(file_path)
                    if ole.exists('WordDocument'):
                        word_stream = ole.openstream('WordDocument').read()
                        # Простое извлечение текста (наивный подход)
                        text = word_stream.decode('latin-1', errors='ignore')
                        # Фильтруем печатные символы
                        text = ''.join(c for c in text if c.isprintable() or c in '\n\r\t')
                        ole.close()
                        if len(text) > 100:
                            self.logger.info(f"Olefile successful | length={len(text)}")
                            return text
                    else:
                        self.logger.warning("WordDocument stream not found")
            except Exception as e:
                self.logger.warning(f"Olefile error | error={type(e).__name__}: {e}")
            
            # Если olefile не сработал, возвращаем ошибку
            return f"ERROR: Cannot parse old .doc format. Please convert to .docx manually."
        
        # Для .docx используем python-docx
        try:
            doc = Document(file_path)
            paragraphs = []
            
            for element in doc.element.body:
                if isinstance(element, CT_P):
                    para = Paragraph(element, doc)
                    if para.text.strip():
                        # Сохраняем заголовки
                        if para.style.name.startswith('Heading'):
                            level = para.style.name.replace('Heading ', '')
                            if level.isdigit():
                                paragraphs.append(f"{'#' * int(level)} {para.text}")
                            else:
                                paragraphs.append(f"**{para.text}**")
                        else:
                            paragraphs.append(para.text)
                            
                elif isinstance(element, CT_Tbl):
                    table = Table(element, doc)
                    paragraphs.append(self._table_to_markdown(table))
            
            return "\n\n".join(paragraphs)
            
        except Exception as e:
            self.logger.error(f"Fallback parser failed | error={type(e).__name__}: {e}")
            return f"ERROR: Failed to parse document: {str(e)}"
    
    def _table_to_markdown(self, table: 'Table') -> str:
        """Конвертация таблицы Word в Markdown"""
        if not table.rows:
            return ""
        
        lines = []
        
        # Заголовок таблицы (первая строка)
        header_cells = [cell.text.strip() for cell in table.rows[0].cells]
        lines.append("| " + " | ".join(header_cells) + " |")
        lines.append("|" + "|".join(["---"] * len(header_cells)) + "|")
        
        # Остальные строки
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
        
        return "\n".join(lines)
    
    def _process_images_with_ocr_individually(self, images: List[Dict]) -> List[str]:
        """
        OCR обработка изображений через Unstructured (возвращает список текстов)
        
        Args:
            images: List изображений с путями
            
        Returns:
            List OCR текстов для каждого изображения (в том же порядке)
        """
        if not images:
            self.logger.info("No images to process with OCR")
            return []
        
        ocr_texts = []
        successful = 0
        failed = 0
        
        self.logger.info(f"Starting OCR for {len(images)} images | strategy={self.ocr_strategy}")
        
        for img in images:
            try:
                # Проверка существования файла
                if not os.path.exists(img['path']):
                    self.logger.error(f"Image file not found | index={img['index']} path={img['path']}")
                    ocr_texts.append("")  # Пустой текст для этого изображения
                    failed += 1
                    continue
                
                # Проверка размера файла
                file_size = os.path.getsize(img['path'])
                if file_size == 0:
                    self.logger.error(f"Image file is empty | index={img['index']} path={img['path']}")
                    ocr_texts.append("")
                    failed += 1
                    continue
                
                self.logger.info(f"Processing image with OCR | index={img['index']} type={img['type']} size={file_size} path={img['path']}")
                
                # Используем partition_image с OCR + русский язык
                elements = partition_image(
                    filename=img['path'],
                    strategy=self.ocr_strategy,
                    infer_table_structure=True,
                    languages=["rus", "eng"]  # КРИТИЧНО: русский + английский
                )
                
                if not elements:
                    self.logger.warning(f"No OCR elements extracted | index={img['index']}")
                    ocr_texts.append("")
                    failed += 1
                    continue
                
                # Извлекаем текст из элементов
                image_text = "\n\n".join([str(el) for el in elements if str(el).strip()])
                
                if image_text.strip():
                    ocr_texts.append(image_text)
                    successful += 1
                    self.logger.info(f"OCR completed | index={img['index']} text_length={len(image_text)}")
                else:
                    self.logger.warning(f"OCR produced empty text | index={img['index']}")
                    ocr_texts.append("")
                    failed += 1
                
            except Exception as e:
                self.logger.error(f"OCR failed | index={img['index']} error={type(e).__name__}: {e}")
                ocr_texts.append("")
                failed += 1
            
            finally:
                # Всегда удаляем временный файл
                try:
                    if os.path.exists(img['path']):
                        os.remove(img['path'])
                        self.logger.debug(f"Cleaned up temp image | path={img['path']}")
                except Exception as e:
                    self.logger.warning(f"Failed to remove temp image | path={img['path']} error={e}")
        
        self.logger.info(f"OCR processing complete | total={len(images)} successful={successful} failed={failed}")
        
        return ocr_texts
    
    def _process_images_with_ocr(self, images: List[Dict]) -> str:
        """
        OCR обработка изображений через Unstructured (возвращает объединенный текст)
        
        Этот метод сохранен для обратной совместимости.
        Использует _process_images_with_ocr_individually внутри.
        
        Args:
            images: List изображений с путями
            
        Returns:
            Объединенный текст из всех изображений
        """
        ocr_texts = self._process_images_with_ocr_individually(images)
        return "\n\n".join(ocr_texts)

