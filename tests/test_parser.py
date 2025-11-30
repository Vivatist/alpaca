"""
Тесты для модулей парсинга документов
"""
import os
import shutil

import pytest

from app.parsers.word_parser_module.word_parser import WordParser
from app.parsers.pptx_parser_module.pptx_parser import PowerPointParser
from app.parsers.excel_parser_module.excel_parser import ExcelParser
from utils.file_manager import File


class TestParserWord:
    """Тесты парсинга DOCX файлов"""
    
    def test_parse_docx_file(self, temp_docx_file):
        """Тест парсинга реального DOCX файла"""
        file = File(hash='test_hash_123', path=temp_docx_file, status_sync='added')
        parser = WordParser(enable_ocr=False)
        
        result = parser.parse(file)
        
        assert result is not None
        assert isinstance(result, str)
        assert len(result) > 0
        assert "Заголовок тестового документа" in result
        assert "первый параграф" in result
    
    def test_parse_nonexistent_file(self):
        """Тест парсинга несуществующего файла"""
        file = File(hash='test_hash_456', path='/nonexistent/file.docx', status_sync='added')

        with pytest.raises(FileNotFoundError):
            WordParser(enable_ocr=False).parse(file)
    
    def test_parse_empty_docx(self):
        """Тест парсинга пустого DOCX файла"""
        from docx import Document
        import tempfile
        
        with tempfile.NamedTemporaryFile(mode='wb', suffix='.docx', delete=False) as f:
            doc = Document()
            doc.save(f.name)
            temp_path = f.name
        
        try:
            file = File(hash='test_hash_empty', path=temp_path, status_sync='added')
            
            with pytest.raises(ValueError):
                WordParser(enable_ocr=False).parse(file)
        finally:
            if os.path.exists(temp_path):
                os.unlink(temp_path)


class TestPowerPointParser:
    """Тесты парсинга PPTX файлов"""

    def test_parse_pptx_file(self, temp_pptx_file):
        file = File(hash='pptx_hash_123', path=temp_pptx_file, status_sync='added')
        parser = PowerPointParser()

        result = parser.parse(file)

        assert isinstance(result, str)
        assert "Тестовая презентация" in result
        assert "Колонка A" in result

    def test_parse_nonexistent_pptx(self):
        file = File(hash='pptx_missing', path='/nonexistent/file.pptx', status_sync='added')
        parser = PowerPointParser()

        with pytest.raises(FileNotFoundError):
            parser.parse(file)
    
    def test_parse_docx_with_multiple_paragraphs(self):
        """Тест парсинга DOCX с несколькими параграфами"""
        from docx import Document
        import tempfile
        
        with tempfile.NamedTemporaryFile(mode='wb', suffix='.docx', delete=False) as f:
            doc = Document()
            paragraphs = [
                "Первый параграф документа.",
                "Второй параграф документа.",
                "Третий параграф документа.",
                "Четвёртый параграф документа.",
            ]
            for p in paragraphs:
                doc.add_paragraph(p)
            doc.save(f.name)
            temp_path = f.name
        
        try:
            file = File(hash='test_hash_multi', path=temp_path, status_sync='added')
            
            result = WordParser(enable_ocr=False).parse(file)
            
            assert result is not None
            assert len(result) > 0
            # Проверяем что все параграфы присутствуют
            for p in paragraphs:
                assert p in result
        finally:
            if os.path.exists(temp_path):
                os.unlink(temp_path)


class TestExcelParser:
    """Тесты парсинга Excel файлов"""

    def test_parse_xlsx_file(self, temp_xlsx_file):
        file = File(hash='xlsx_hash', path=temp_xlsx_file, status_sync='added')
        parser = ExcelParser(max_rows_per_table=50)

        result = parser.parse(file)

        assert isinstance(result, str)
        assert "Лист: Сводка" in result
        assert "Доход" in result
        assert "1250000.75" in result or "1250000.75".rstrip("0") in result
        assert "Лист: Детализация" in result
        assert "Бурение" in result

    def test_parse_nonexistent_xlsx(self):
        file = File(hash='missing_xlsx', path='/tmp/absent.xlsx', status_sync='added')
        parser = ExcelParser()

        with pytest.raises(FileNotFoundError):
            parser.parse(file)

    def test_parse_xls_triggers_conversion(self, temp_xlsx_file, tmp_path, monkeypatch):
        legacy_path = tmp_path / "legacy.xls"
        shutil.copy(temp_xlsx_file, legacy_path)

        called = {"value": False}

        def fake_convert(path):
            assert str(path) == str(legacy_path)
            called["value"] = True
            return temp_xlsx_file

        monkeypatch.setattr(
            "app.parsers.excel_parser_module.excel_parser.convert_xls_to_xlsx",
            fake_convert,
        )

        file = File(hash='xls_hash', path=str(legacy_path), status_sync='added')
        parser = ExcelParser()

        result = parser.parse(file)

        assert called["value"] is True
        assert isinstance(result, str)

    def test_multiline_cells_render_single_line(self):
        from openpyxl import Workbook
        import tempfile
        from pathlib import Path

        with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as tmp:
            workbook_path = Path(tmp.name)

        try:
            wb = Workbook()
            sheet = wb.active
            sheet.append(["Колонка A", "Колонка B"])
            sheet.append(["Строка\n1", "Text\nwith\nlinebreaks"])
            wb.save(workbook_path)

            file = File(hash='multi_hash', path=str(workbook_path), status_sync='added')
            parser = ExcelParser()

            result = parser.parse(file)

            assert "Строка 1" in result
            assert "Text with linebreaks" in result
            assert "Строка\n1" not in result
        finally:
            if workbook_path.exists():
                workbook_path.unlink()

    def test_parse_xls_without_conversion(self, temp_xls_file, monkeypatch):
        convert_calls = {"value": 0}

        def fail_convert(path):
            convert_calls["value"] += 1
            return None

        monkeypatch.setattr(
            "app.parsers.excel_parser_module.excel_parser.convert_xls_to_xlsx",
            fail_convert,
        )

        file = File(hash='xls_fallback_hash', path=str(temp_xls_file), status_sync='added')
        parser = ExcelParser()

        result = parser.parse(file)

        assert convert_calls["value"] == 1
        assert "Лист: План" in result
        assert "Старт" in result
        assert "2025-01-05" in result
        assert "750000" in result
