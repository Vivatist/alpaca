#!/usr/bin/env python3
"""
Metadata Extractor - извлечение метаданных из Word документов

Извлекает специфичные для Word метаданные (автор, страницы, параграфы и т.д.)
"""

from typing import Dict
from docx import Document  # type: ignore

from logging_config import get_logger

logger = get_logger("core.parser.metadata_extractor")


def extract_word_metadata(file_path: str) -> Dict:
    """
    Извлечение СПЕЦИФИЧНЫХ для Word метаданных
    
    Общие метаданные (file_name, file_path, file_size, etc.) добавляются в базовом классе.
    Здесь добавляем только специфичные для Word данные.
    
    Args:
        file_path: Путь к DOCX файлу
        
    Returns:
        Dict со специфичными метаданными (author, subject, pages, paragraphs, tables, images)
    """
    specific_metadata = {
        'author': '',
        'subject': '',
        'pages': 0,
        'paragraphs': 0,
        'tables': 0,
        'images': 0
    }
    
    try:
        doc = Document(file_path)
        core_props = doc.core_properties
        
        specific_metadata['author'] = core_props.author or ''
        specific_metadata['subject'] = core_props.subject or ''
        
        # Подсчет элементов
        specific_metadata['paragraphs'] = len(doc.paragraphs)
        specific_metadata['tables'] = len(doc.tables)
        
        # Приблизительное количество страниц (250 слов на страницу)
        total_words = sum(len(p.text.split()) for p in doc.paragraphs)
        specific_metadata['pages'] = max(1, total_words // 250)
        
        # Подсчет изображений
        image_count = 0
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_count += 1
        specific_metadata['images'] = image_count
        
        logger.debug(f"Word-specific metadata | author={specific_metadata['author']} pages={specific_metadata['pages']} paragraphs={specific_metadata['paragraphs']}")
        
    except Exception as e:
        logger.warning(f"Failed to extract Word-specific metadata | file={file_path} error={e}")
    
    return specific_metadata
